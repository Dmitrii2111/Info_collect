from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "designer_brief.md"
OUTPUT_PATH = BASE_DIR / "designer_brief.docx"


def add_markdown_to_doc(document: Document, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line[:2].isdigit() and line[1:3] == ". ":
            document.add_paragraph(line[3:].strip(), style="List Number")
            continue

        document.add_paragraph(line)


def main() -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    markdown_text = SOURCE_PATH.read_text(encoding="utf-8")
    add_markdown_to_doc(document, markdown_text)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
